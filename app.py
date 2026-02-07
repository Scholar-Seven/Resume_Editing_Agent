import os
import uuid
from datetime import datetime
from typing import List, Optional, Dict, Any

from flask import Flask, request, jsonify, send_file, render_template
from pydantic import BaseModel, Field, ValidationError

import pdfplumber
from docx import Document

from langchain_deepseek import ChatDeepSeek
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser

import re
from docx.shared import Pt
from docx.oxml.ns import qn

# ----------------------------
# Config
# ----------------------------
UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
if not DEEPSEEK_API_KEY:
    raise RuntimeError("Please set env DEEPSEEK_API_KEY")

# 使用 deepseek-chat（支持结构化输出/工具能力；reasoner 不支持这些特性）
llm = ChatDeepSeek(
    model="deepseek-chat",
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com",
    temperature=0.2,
)

app = Flask(__name__)

# ----------------------------
# In-memory store (demo)
# Production: use SQLite/Redis
# ----------------------------
DB: Dict[str, Dict[str, Any]] = {}


# ----------------------------
# Utils: file -> text
# ----------------------------
def extract_text_from_pdf(path: str) -> str:
    texts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t.strip():
                texts.append(t)
    return "\n".join(texts).strip()

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(parts).strip()

def file_to_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    if ext in [".docx", ".doc"]:
        return extract_text_from_docx(path)
    raise ValueError(f"Unsupported file type: {ext}")

def to_dict(obj):
    if isinstance(obj, dict):
        return obj
    if hasattr(obj, "model_dump"):  # pydantic v2
        return obj.model_dump()
    if hasattr(obj, "dict"):        # pydantic v1
        return obj.dict()
    return obj

# ----------------------------
# Pydantic schemas
# ----------------------------
class ExperienceBullet(BaseModel):
    text: str
    tags: List[str] = Field(default_factory=list)

class WorkExperience(BaseModel):
    company: str = ""
    title: str = ""
    start: str = ""
    end: str = ""
    bullets: List[ExperienceBullet] = Field(default_factory=list)

class ResumeProfile(BaseModel):
    name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    summary: str = ""
    skills: List[str] = Field(default_factory=list)
    work_experiences: List[WorkExperience] = Field(default_factory=list)
    education: List[str] = Field(default_factory=list)
    projects: List[str] = Field(default_factory=list)
    raw_keywords: List[str] = Field(default_factory=list)

class Suggestion(BaseModel):
    id: str
    priority: str  # P0/P1/P2
    target_section: str
    issue: str
    recommendation: str
    rewrite_hint: str = ""
    evidence: str = ""

class MatchReport(BaseModel):
    overall_score: int  # 0-100
    strengths: List[str] = Field(default_factory=list)
    gaps: List[str] = Field(default_factory=list)
    suggestions: List[Suggestion] = Field(default_factory=list)


# ----------------------------
# Chains
# ----------------------------
def parse_resume_chain(resume_text: str) -> ResumeProfile:
    parser = JsonOutputParser(pydantic_object=ResumeProfile)

    prompt = ChatPromptTemplate.from_messages([
        ("system",
         "你是资深HR与简历结构化专家。"
         "请把用户简历内容抽取为结构化JSON。"
         "只基于文本，不要编造。输出必须严格符合JSON schema。"),
        ("user",
         "简历原文如下：\n\n{resume_text}\n\n{format_instructions}")
    ])

    chain = prompt | llm | parser
    return chain.invoke({
        "resume_text": resume_text,
        "format_instructions": parser.get_format_instructions()
    })

def analyze_match_chain(profile: ResumeProfile, jd_text: str) -> MatchReport:
    parser = JsonOutputParser(pydantic_object=MatchReport)

    prompt = ChatPromptTemplate.from_messages([
        ("system",
         "你是招聘经理+简历优化顾问。"
         "任务：对比简历与JD，给出匹配评分、优势、缺口，并输出可执行优化建议。"
         "约束：建议必须具体到段落/要点；不要建议造假；缺少信息就建议用户补充占位符。"
         "输出必须是严格JSON。"),
        ("user",
         "【简历结构化信息】\n{profile_json}\n\n【JD】\n{jd_text}\n\n{format_instructions}")
    ])

    chain = prompt | llm | parser
    return chain.invoke({
        "profile_json": profile if isinstance(profile, dict) else profile.dict(),
        "jd_text": jd_text,
        "format_instructions": parser.get_format_instructions()
    })

def rewrite_resume_chain(resume_text: str, jd_text: str, confirmed_suggestions: List[Suggestion]) -> str:
    # 这里直接让模型输出“优化后的简历纯文本”（再写入 docx）
    # 你也可以让它输出分段JSON，再更稳定地落盘
    prompt = ChatPromptTemplate.from_messages([
        ("system",
         "你是专业简历改写助手。"
         "只允许基于原简历事实进行重写与表达优化，不得新增不存在的经历/公司/项目/数据。"
         "如果建议需要数据但简历没有，请用【待补充：...】占位符。"
         "输出为中文简历正文，结构清晰，适配JD关键词。"),
        ("user",
         "【JD】\n{jd_text}\n\n"
         "【原简历】\n{resume_text}\n\n"
         "【用户已确认的优化建议】\n{suggestions}\n\n"
         "请根据确认建议生成“优化后的完整简历文本”。")
    ])

    suggestions_text = "\n".join(
        [f"- ({s.priority})[{s.target_section}] {s.recommendation} | hint: {s.rewrite_hint}"
         for s in confirmed_suggestions]
    )

    resp = (prompt | llm).invoke({
        "jd_text": jd_text,
        "resume_text": resume_text,
        "suggestions": suggestions_text
    })
    return resp.content


# ----------------------------
# Output: write docx
# ----------------------------
def write_docx(text: str, out_path: str):
    """
    规则：
    - 中文：仿宋，五号（10.5pt）
    - 英文：Times New Roman，五号（10.5pt）
    """
    doc = Document()

    # 中文正则
    zh_pattern = re.compile(r'[\u4e00-\u9fff]')

    for line in text.splitlines():
        p = doc.add_paragraph()

        # 按“中/非中”拆分字符
        buffer = ""
        buffer_is_zh = None

        def flush_buffer(buf, is_zh):
            if not buf:
                return
            run = p.add_run(buf)
            run.font.size = Pt(10.5)  # 五号

            if is_zh:
                # 中文：仿宋
                run.font.name = "仿宋"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            else:
                # 英文：Times New Roman
                run.font.name = "Times New Roman"

        for ch in line:
            is_zh = bool(zh_pattern.search(ch))
            if buffer_is_zh is None:
                buffer = ch
                buffer_is_zh = is_zh
            elif is_zh == buffer_is_zh:
                buffer += ch
            else:
                flush_buffer(buffer, buffer_is_zh)
                buffer = ch
                buffer_is_zh = is_zh

        flush_buffer(buffer, buffer_is_zh)

    doc.save(out_path)



# ----------------------------
# Routes
# ----------------------------
@app.get("/")
def index():
    return render_template("index.html")


@app.post("/upload")
def upload():
    """
    form-data:
      - resume_file: file
      - jd_text: str
    """
    f = request.files.get("resume_file")
    jd_text = request.form.get("jd_text", "").strip()
    if not f or not jd_text:
        return jsonify({"error": "resume_file and jd_text are required"}), 400

    resume_id = str(uuid.uuid4())
    filename = f"{resume_id}_{f.filename}"
    path = os.path.join(UPLOAD_DIR, filename)
    f.save(path)

    DB[resume_id] = {
        "resume_path": path,
        "jd_text": jd_text,
        "resume_text": None,
        "profile": None,
        "report": None,
        "confirmed_suggestion_ids": [],
        "created_at": datetime.utcnow().isoformat()
    }
    return jsonify({"resume_id": resume_id})

@app.post("/analyze")
def analyze():
    data = request.get_json(force=True)
    resume_id = data.get("resume_id", "")
    if resume_id not in DB:
        return jsonify({"error": "resume_id not found"}), 404

    rec = DB[resume_id]
    if not rec["resume_text"]:
        rec["resume_text"] = file_to_text(rec["resume_path"])

    try:
        profile = parse_resume_chain(rec["resume_text"])
        report = analyze_match_chain(profile, rec["jd_text"])
    except ValidationError as e:
        return jsonify({"error": "LLM output validation failed", "detail": str(e)}), 500

    rec["profile"] = to_dict(profile)
    rec["report"] = to_dict(report)
    return jsonify(rec["report"])

@app.post("/confirm")
def confirm():
    """
    json:
      - resume_id
      - confirmed_suggestion_ids: [..]  # 用户确认要应用的建议id
    """
    data = request.get_json(force=True)
    resume_id = data.get("resume_id", "")
    ids = data.get("confirmed_suggestion_ids", [])
    if resume_id not in DB:
        return jsonify({"error": "resume_id not found"}), 404
    if not isinstance(ids, list):
        return jsonify({"error": "confirmed_suggestion_ids must be list"}), 400

    DB[resume_id]["confirmed_suggestion_ids"] = ids
    return jsonify({"ok": True, "confirmed": ids})

@app.post("/optimize")
def optimize():
    data = request.get_json(force=True)
    resume_id = data.get("resume_id", "")
    if resume_id not in DB:
        return jsonify({"error": "resume_id not found"}), 404

    rec = DB[resume_id]
    if not rec["report"]:
        return jsonify({"error": "please run /analyze first"}), 400

    report = MatchReport(**rec["report"])
    confirmed = set(rec["confirmed_suggestion_ids"])
    confirmed_suggestions = [s for s in report.suggestions if s.id in confirmed]

    if not confirmed_suggestions:
        return jsonify({"error": "no confirmed suggestions"}), 400

    optimized_text = rewrite_resume_chain(rec["resume_text"], rec["jd_text"], confirmed_suggestions)

    out_name = f"{resume_id}_optimized.docx"
    out_path = os.path.join(OUTPUT_DIR, out_name)
    write_docx(optimized_text, out_path)

    rec["output_path"] = out_path
    return jsonify({"ok": True, "download_url": f"/download/{resume_id}"})

@app.get("/download/<resume_id>")
def download(resume_id: str):
    rec = DB.get(resume_id)
    if not rec or "output_path" not in rec:
        return jsonify({"error": "file not ready"}), 404
    return send_file(rec["output_path"], as_attachment=True)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
